import os
import docx
os.chdir(r"C:\\FileTransferFTP\\eDoc")
out_folder = os.curdir
print os.path.abspath(out_folder)

for i in range(5000):
    document = docx.Document()
    paragraph = document.add_paragraph('This is a demo Document')
    document.add_heading('This document is index number {}'.format(i))
    document.save('DemoDocument{}.docx'.format(i))
